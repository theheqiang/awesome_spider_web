# !/usr/bin/env python
# -*- coding: utf-8 -*-

import re
import os
from threading import Thread
from multiprocessing import JoinableQueue
from io import BytesIO

import requests
from lxml import etree
from lxml.html import tostring
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import  Pt
from docx.shared import RGBColor



def clear_string(source_list, string):
    '''
        清除列表中存在的特定字符串
        :param source_list: 要处理的列表
        :param string :要清除的字符串
    '''
    dist_list = []
    for item in source_list:
        if string in item:
            # print()
            dist_list.append(re.sub(string, '', item))
        else:
            dist_list.append(item)
    return dist_list


def join_my_answer(type, my_answer):
    '''
    由于判断题的答案在列表中是分开的，所以要将列表中判断题的答案连接在一起
    :param type: 题目类型
    :param my_answer: 答案
    :return: 拼接好的答案
    '''
    try:
        index = type.index('判断题')
        answer = my_answer[0:index]
        for i in range(index, len(my_answer), 2):
            a = ''.join(my_answer[i:i + 2])
            answer.append(a)
        return answer
    except ValueError:
        return my_answer


def deal_answer(answer):
    '''
    对答案进行处理，将其替换成'√', '×'
    :param answer:
    :return:
    '''
    dist_answer = []
    for judge in answer:
        if judge == 'fr dui':
            dist_answer.append('√')
        else:
            dist_answer.append('×')
    return dist_answer


def comb_question(type, question, items, select_items, my_answer, judge_answer):
    '''
    将问题的所有项进行组合，
    :param type: 题目类型
    :param question: 题目
    :param items: 选项（选择题）
    :param select_items: 选择内容
    :param my_answer: 我的答案
    :param judge_answer: 对我的答案的判断
    :return: 组合好的问题
    '''
    # 先将 题目类型、题目、我的答案、答案判断 组合成一个元组，放在列表里面
    comb_ok_question = list(zip(type, question, my_answer, judge_answer))
    # 然后把选择题的选项组合好
    i = 0
    # 由于这里的选择题只有四个选项
    # TODO 由于有的学校选择题不止四个答案， 四个答案的用range(0, len(items), 4)， 五个答案的用range(0, len(items), 5)，（对于全部选择题一样）
    # 如果选择题个数不一样的话，建议从items中入手，或者在内容提取时作出相应处理，这里就不作处理了
    for index in range(0, len(items), 4):
        options = tuple(zip(items[index:index + 4], select_items[index:index + 4]))
        comb_ok_question[i] += tuple(options)
        i += 1
    return comb_ok_question


def getcookies(base_url, header):
    '''
    获取登录该网站的cookie
    :param base_url: 该网站的起始网址
    :return: cookie
    '''


    # 存储cookie
    all_cookie = {}
    try:
        # 判断本地是否有cookie.txt文件
        if not os.path.exists('cookie.txt'):
            # 请求超星网址
            base_response = requests.get(
                url=base_url,
                headers=header,
                timeout=10
            )

            # 更新cookie
            all_cookie.update(base_response.cookies.get_dict())

            # uuid = re.findall(r'<input type = "hidden" value="(.*)" id = "uuid"/>', base_response.text)
            # enc = re.findall(r'<input type = "hidden" value="(.*)" id = "enc"/>', base_response.text)
            # quickCode = re.findall(r' <img src="(.*)" id="quickCode">', base_response.text)

            tree = etree.HTML(base_response.text)
            uuid = tree.xpath('//input[@id="uuid"]/@value')
            enc = tree.xpath('//input[@id="enc"]/@value')
            quickCode = tree.xpath('//img[@id="quickCode"]/@src')

            code_url = 'https://passport2.chaoxing.com'
            # with open("./code.jpg", "wb")as f:
            #     f.write(requests.get(
            #         url=code_url + quickCode[0],
            #         timeout=10
            #     ).content)
            # 显示二维码
            page_resource = requests.get(url=code_url + quickCode[0],timeout=10).content

            im = Image.open(BytesIO(page_resource))
            im.show()

            # 等待扫码
            state = input('扫码完成？ Y/N\n')
            if state is 'Y':
                # 1. 扫完码请求登录链接
                passport_url = 'https://passport2.chaoxing.com/getauthstatus'  # 更新header
                header['Accept'] = 'application/json, text/javascript, */*; q=0.01'
                header['Host'] = 'passport2.chaoxing.com'
                header['Origin'] = 'https://passport2.chaoxing.com'
                header[
                    'Referer'] = 'https://passport2.chaoxing.com/login?fid=&newversion=true&refer=http%3A%2F%2Fi.chaoxing.com'

                data = {
                    'enc': enc,
                    'uuid': uuid
                }

                passport_response = requests.post(
                    url=passport_url,
                    headers=header,
                    data=data,
                    cookies=all_cookie,
                    timeout=10
                )
            # 更新cookie
            all_cookie.update(passport_response.cookies.get_dict())
            # 可以在这里将cookie信息保存，在以后的运行中就不用扫码了
            f = open("cookie.txt", 'w')
            f.write(str(all_cookie))
            f.close()
        else:
            # 本地已经有cookie.txt, 读取cookie
            f = open("cookie.txt", 'r')
            all_cookie = eval(f.read())
            f.close()
    except:
        # 请求超时
        print('请求cookie文件超时')
    return all_cookie


def getquestions(queue, base_url, all_cookie, header):
    # 章节名称， 用于存储有题目的章节的名称
    # 设置超时

    chapter_name_list = []

    header[
        'Accept'] = 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3'
    header['Host'] = 'i.chaoxing.com'
    # 请求下面这个链接，自动重定向到url：http://i.chaoxing.com/base?t=1589036774406 获取课程选择页面参数
    base_response = requests.get(
        url=base_url,
        headers=header,
        cookies=all_cookie,
        timeout=10
    )
    all_cookie.update(base_response.cookies.get_dict())
    # param: interaction?s=f94269a49525d8a229c414861fbbd65c
    param = re.findall(r"(.*)visit/(.*)',this", base_response.text)[0][1]

    visit_url = 'http://mooc1-1.chaoxing.com/visit/' + param
    header['Host'] = 'mooc1-1.chaoxing.com'
    header['Referer'] = 'http://i.chaoxing.com/base?t=1588997030446'
    # 请求课程选择页面
    visit_response = requests.get(
        url=visit_url,
        headers=header,
        cookies=all_cookie,
        timeout=10
    )
    all_cookie.update(visit_response.cookies.get_dict())

    # 获取第一个课程的参数,如果要获取所有课程，在这里使用for,注意：这里两个链接一课课程，例如第一个课程[0][1],[1][1]都可以
    # 虽然他们的不一样的，但都能打开，同理，第二个课程[2][1], [3][1]都可以。
    # param = re.findall(r"(.*)/mycourse(.*)'", visit_response.text)[0][1]
    # course_url = 'https://mooc1-1.chaoxing.com/mycourse' + param
    course_dict = {}
    tree = etree.HTML(visit_response.text)
    li_list = tree.xpath('/html/body/div/div[2]/div[3]/ul/li')
    li_list.pop()
    print('序号','    课程名称')
    for index,li in enumerate(li_list,start=1):
        course_url = 'https://mooc1-1.chaoxing.com' + li.xpath('./div/a/@href')[0]
        course_title = li.xpath('./div[2]/h3/a/@title')[0]
        course_dict[index] = {'url':course_url,'title':course_title}
        print(index,'    '+course_title)

    while 1:
        chose_course = input('请选择你要保存课程章节的序号(q退出):').strip()
        if chose_course.upper() == 'Q':
            break
        try:
            course_url = course_dict.get(int(chose_course)).get('url')
        except:
            print('输入的序号有误,请重新输入')
            continue
        else:
        # 在课程url中匹配出所有的id信息
            params = re.findall(r".*courseId=(.*)&clazzid=(.*)&vc=(.*)&cpi=(.*)&enc=(.*)", course_url)

            courseId = params[0][0]
            clazzid = params[0][1]
            vc = params[0][2]
            cpi = params[0][3]
            chapter_enc = params[0][4]
            # 请求第一个课程的链接, 得到所有章节的信息
            course_response = requests.get(
                url=course_url,
                headers=header,
                cookies=all_cookie,
                timeout=10
            )

            # 这里直接用正则表达式太慢了，先用xpath匹配出所有的章节的span信息
            selector = etree.HTML(course_response.text)
            # 提取课程名称
            course_name = course_dict.get(int(chose_course)).get('title')

            for ii in selector.xpath("//span[@class='articlename']/a"):

                # 将byte类型转换为string类型
                s = bytes.decode(tostring(ii))
                chapterId = re.findall(r'.*chapterId=(.*)&amp;courseId', s)[0]
                # 组建章节的url
                chapter_url = 'https://mooc1-1.chaoxing.com/mycourse/studentstudy?chapterId={}&courseId={}&clazzid={}&enc={}'.format(
                    chapterId, courseId, clazzid, chapter_enc)

                chapter_response = requests.get(
                    url=chapter_url,
                    headers=header,
                    cookies=all_cookie,
                    timeout=10
                )
                utenc = re.findall(r'.*utEnc="(.*)"', chapter_response.text)[0]
                # 请求该url获取该章节的标题
                studentstudyAjax_url = 'https://mooc1-1.chaoxing.com/mycourse/studentstudyAjax'
                data = {
                    'courseId': courseId,
                    'clazzid': clazzid,
                    'chapterId': chapterId,
                    'cpi': cpi,
                    'verificationcode': ''
                }
                studentstudyAjax_response = requests.post(
                    url=studentstudyAjax_url,
                    headers=header,
                    data=data,
                    cookies=all_cookie,
                    timeout=10
                )
                # 章节标题
                chapter_name = re.findall(r'.*<h1>(.*)</h1>', studentstudyAjax_response.text)[0]

                url = 'https://mooc1-1.chaoxing.com/knowledge/cards?' \
                      'clazzid={}&courseid={}&knowledgeid={}&num=1&ut=s&cpi={}&v=20160407-1'.format(clazzid, courseId,
                                                                                                    chapterId, cpi)

                # TODO 上面的num=取决于你想保存哪种类型的作业题,有两种,一种是章节测验,另一种是作业,具体看自己的需求要选择对应的num值
                header['Referer'] = course_url
                cards_respose = requests.get(
                    url=url,
                    headers=header,
                    cookies=all_cookie,
                    timeout=10
                )

                # 没有题目的章节在这里会报错
                try:
                    workId = re.findall(r'"jobid":"work-(.*?)"', cards_respose.text)[0]
                    # 有题目
                    print(chapter_name + '有题目')
                    chapter_name_list.append(chapter_name)
                except:
                    print(chapter_name + '没有题目')
                    continue
                enc = re.findall(r'.*"enc":"(.*)","type":"workid', cards_respose.text)[0]
                # TODO 查找出下面所需参数 ktoken
                ktoken = re.findall(r'.*"ktoken":"(.*)","isFiled', cards_respose.text)[0]
                header[
                    'Referer'] = 'https://mooc1-1.chaoxing.com/ananas/modules/work/index.html?v=2020-0612-1728&castscreen=0'
                params = (
                    ('api', vc),
                    ('workId', workId),
                    ('jobid', 'work-{}'.format(workId)),
                    ('needRedirect', 'true'),
                    ('knowledgeid', chapterId),
                    ('ktoken', ktoken),
                    ('cpi', cpi),
                    ('ut', 's'),
                    ('clazzId', clazzid),
                    ('type', ''),
                    ('enc', enc),
                    ('utenc', utenc),
                    ('courseid', courseId),
                )

                # # TODO 这里加了两个参数， ktoken和cpi， cpi的值和上面的是一样的
                # param = 'api={}&workId={}&jobid={}&needRedirect={}&knowledgeid={}&ktoken={}&ut={}&clazzId={}&type={}&cpi={}&enc={}&utenc={}&courseid={}' \
                #     .format(data['api'], data['workId'], data['jobid'], data['needRedirect'], data['knowledgeid'],
                #             data['ktoken'],
                #             data['ut'], data['clazzId'], data['type'], cpi
                #             , data['enc'], data['utenc'], data['courseid'])
                # # 这里不能用data=data， 要用上面的这个组装好的参数，还不知道为什么
                work_response = requests.get(
                    url='https://mooc1-1.chaoxing.com/api/work',
                    params = params,
                    headers=header,
                    cookies=all_cookie,
                    timeout=10
                )
                # 将网页内的内容生成一个树
                selector = etree.HTML(work_response.text)
                question_text = ''.join(selector.xpath("//div[@style='width:80%;height:100%;float:left;']//text()"))

                # 用正则表达式提取题目类型
                com = re.compile(r'.*【(.*)】')
                types = re.findall(com, question_text)

                # 去掉文本中的制表符、换行符、空格
                question_text = re.sub(r'(\t|\n|\s)?', '', question_text)

                # 提取题目，第一个为空，去掉第一个
                question = re.split('【[\u4e00-\u9fa5]{3}】', question_text)[1:]
                # 选项， A B C D
                items = selector.xpath("//li[@class='clearfix']/i//text()")

                # 网页中提取的选项可能含有空格等数据，要进行清除 ；选项内容
                select_items = clear_string(selector.xpath("//li[@class='clearfix']/a//text()"), '\xa0')

                # 我的答案
                # my_answer = join_my_answer(types, selector.xpath("//div[@class='Py_answer clearfix']/span//text()"))# 移动平台用这个
                my_answer = join_my_answer(types, selector.xpath("//div[@class='Py_answer clearfix']/span[1]//text()"))
                # 我的答案的判断
                judge_answer = deal_answer(selector.xpath("//div[@class='Py_answer clearfix']/i/@class"))
                # 对上面处理好的问题项进行组合, 然后放进队列中
                comb_ok_question = comb_question(types, question, items, select_items, my_answer, judge_answer)
                course_chapter = tuple([course_name, chapter_name])
                comb_ok_question.append(course_chapter)
                queue.put(comb_ok_question)

    queue.join()

def qustion_to_word(queue, document):
    # 设置一个空白样式
    style = document.styles['Normal']
    # 设置西文字体
    style.font.name = 'Times New Roman'
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    coursenames = set()
    # 插入课程名称标识
    flag = True
    # 设置超时退出，由于我这个课程有些课的章节还没有开发，设置长一点,3分钟

    while True:
        # 由于有些章节（并且数量还很多）会没开放，程序会陷入很长的等待中，所以在这里设置超时
        try:
            question = queue.get(block=True, timeout=120)
        except:
            return
        # question = [('单选题', 'Dalvik虚拟机是基于（）的架构。', '我的答案：C', '×', ('A、', '栈'), ('B、', '堆'), ('C、', '寄存器'), ('D、', '存储器')), ('单选题', '关于Dalvik虚拟机，说法错误的是().', '我的答案：B', '√', ('A、', '基于寄存器的架构'), ('B、', '基于栈的架构'), ('C、', '加载的是.dex格式的数据'), ('D、', '在linux操作系统上运行')), ('判断题', '第四代移动通讯技术(4G)包括TD-LTE和FDD-LTE两种制式。', '我的答案：√', '√'), ('判断题', '随着智能手机的发展，移动通信技术也在不断地升级，目前应用最广泛的是4G', '我的答案：√', '√'), ('判断题', 'Android系统采用分层架构，由高到低分为4层，依次是应用程序层、应用程序框架层、核心类库和Linux内核。', '我的答案：√', '√'), ('判断题', 'Android是Google公司基于Linux平台开发的手机及平板电脑操作系统', '我的答案：√', '√'), ('判断题', 'Dalvik虚拟机是在linux操作系统上运行.', '我的答案：√', '√'), ('判断题', 'Android系统最初由安迪·鲁宾等人开发制作。', '我的答案：√', '√'), ('判断题', 'Android是Google公司基于Linux平台开发的手机及平板电脑操作系统。', '我的答案：√', '√'), ('判断题', 'Android底层是基于linux操作系统的。', '我的答案：√', '√'), ('判断题', '所有的Android应用程序在底层对应同一个Dalvik虚拟机实例，其代码在虚拟机的解析得以执行。', '我的答案：×', '√'), ('移动平台开发', 'Android智能手机开发概况')]

        course_name = question[-1][0]
        # 第一个章节的问题
        if (flag or (course_name not in coursenames)):
            # 如果当前章节的课程名称不等于上一个章节的课程，表明这是新的一课；或者是第一课的第一章节
            try:
                # 空集合pop会报错，要过滤掉
                coursenames.pop()
            except:
                pass
            coursenames.add(course_name)
            flag = False
            # 插入插入课程名称
            print('插入插入课程名称')
            # 设置标题
            title = document.add_heading(level=0)
            # 标题居中
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置标题内容
            title_run = title.add_run(course_name)
            # 设置标题字体大小
            title_run.font.size = Pt(14)

        # 首先插入章节标题
        chapter_name = document.add_heading(question[-1][1], level=1)

        # 添加题目, 最后一个是课程名称和章节名称，不要
        # ('单选题', 'Dalvik虚拟机是基于（）的架构。', '我的答案：C', '×', ('A、', '栈'), ('B、', '堆'), ('C、', '寄存器'), ('D、', '存储器'))
        for qes in question[0:len(question) - 1]:
            print(qes)
            # 单选题, 添加题目类型和题目
            type_and_question = qes[0] + ' : ' + qes[1]
            document.add_paragraph(type_and_question, style='List Number')
            if qes[0] != '判断题':
                # 添加选项
                for index in range(4, len(qes)):
                    select_item = qes[index]

                    item_string = select_item[0] + select_item[1]

                    document.add_paragraph(item_string)

            # 添加我的答案和判断
            answer_and_judge = document.add_paragraph()
            aj = qes[2] + ' ' * 50 + qes[3]

            answer_and_judge_run = answer_and_judge.add_run(aj)
            if qes[3] == '√':
                # 添加并设置为绿色
                answer_and_judge_run.font.color.rgb = RGBColor(0, 255, 0)
            else:
                # 设置为红色
                answer_and_judge_run.font.color.rgb = RGBColor(255, 0, 0)

        queue.task_done()


def main():
    queue = JoinableQueue()
    document = Document()
    # 起始网址
    base_url = 'http://i.chaoxing.com'
    # 请求头
    header = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/73.0.3683.103 Safari/537.36'
    }
    try:
        # 获取cookie
        all_cookie = getcookies(base_url, header)

        question_produce = Thread(target=getquestions, args=[queue, base_url, all_cookie, header])
        question_consumer = Thread(target=qustion_to_word, args=[queue, document])

        question_consumer.daemon = True


        question_produce.start()
        question_consumer.start()
        question_produce.join(5 * 60)
    except:
        pass

    finally:
        document.save('学习通课程题目和答案.docx')


if __name__ == '__main__':
    main()